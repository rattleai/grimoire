#!/usr/bin/env python3
"""Inventory existing technical-documentation manuals.

Reads N input files (PDF, DOCX, plain text) from a directory and produces an
inventory JSON: identification, chapter map, coverage matrix, and reusability
candidates.

Usage:
    python inventory_techdocs.py <input_dir> [--out inventory.json]
                                              [--lang de]
                                              [--max-files 20]

The output contract matches the ``inventory.json`` shape documented in
``rattle-techdoc/SKILL.md`` Step 1.

Pure stdlib + ``pypdfium2`` + ``python-docx`` (both optional). If ``pypdfium2``
is missing, PDF files are skipped with a warning.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from dataclasses import asdict, dataclass, field
from pathlib import Path

# ---------------------------------------------------------------------------
# Canonical chapter slugs (mirrors rattle-techdoc/references/chapter-reference.md)
# ---------------------------------------------------------------------------

CANONICAL_CHAPTERS: list[tuple[str, list[str]]] = [
    ("ch-00-cover", ["deckblatt", "cover", "title page"]),
    ("ch-00-toc", ["inhaltsverzeichnis", "table of contents", "contents"]),
    ("ch-01-about-document", ["zu diesem dokument", "about this document", "über dieses dokument"]),
    ("ch-02-safety", ["sicherheit", "safety", "sicherheitshinweise"]),
    (
        "ch-03-product-description",
        [
            "produktbeschreibung",
            "produkt- und systembeschreibung",
            "product description",
            "system description",
        ],
    ),
    ("ch-04-transport", ["transport", "anlieferung", "lagerung", "delivery", "storage"]),
    ("ch-05-assembly", ["montage", "installation", "assembly", "aufstellung"]),
    ("ch-06-commissioning", ["inbetriebnahme", "commissioning", "einstellungen", "settings"]),
    ("ch-07-operation", ["bedienung", "betrieb", "operation"]),
    ("ch-08-troubleshooting", ["störung", "fehlerbehebung", "troubleshooting", "diagnose"]),
    (
        "ch-09-maintenance",
        ["wartung", "reinigung", "maintenance", "cleaning", "inspection", "reparatur", "repair"],
    ),
    (
        "ch-10-modifications",
        ["umbau", "erweiterung", "modernisierung", "modification", "extension", "modernisation"],
    ),
    (
        "ch-11-decommissioning",
        [
            "außerbetriebnahme",
            "demontage",
            "entsorgung",
            "decommissioning",
            "disassembly",
            "disposal",
        ],
    ),
    ("ch-12-conformity", ["konformität", "conformity", "rechtliche hinweise", "legal", "normen"]),
    ("ch-13-appendix", ["anhang", "appendix", "anlagen"]),
]

HEADING_RE = re.compile(
    r"^\s*(?:\d+(?:\.\d+)*\.?\s+)?([A-ZÄÖÜ][^\n]{2,80})$",
    re.MULTILINE,
)


@dataclass
class FileInventory:
    path: str
    bytes_size: int
    extension: str
    detected_language: str = ""
    detected_chapters: list[str] = field(default_factory=list)
    coverage: dict[str, str] = field(default_factory=dict)
    text_excerpt: str = ""


@dataclass
class Inventory:
    input_dir: str
    file_count: int
    total_bytes: int
    files: list[FileInventory]
    canonical_chapters: list[str]
    cross_file_chapter_frequency: dict[str, int]
    reusability_candidates: list[dict]


# ---------------------------------------------------------------------------
# File text extraction
# ---------------------------------------------------------------------------


def _extract_text(path: Path) -> tuple[str, str]:
    """Return (text, detected_extension_used)."""
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore"), ext
    if ext == ".pdf":
        try:
            import pypdfium2 as pdfium  # type: ignore
        except ImportError:
            print(f"[skip] pypdfium2 missing, cannot read {path.name}", file=sys.stderr)
            return "", ext
        pdf = pdfium.PdfDocument(str(path))
        out = []
        for i in range(len(pdf)):
            page = pdf[i]
            out.append(page.get_textpage().get_text_range())
        return "\n".join(out), ext
    if ext in {".docx", ".doc"}:
        try:
            import docx  # type: ignore
        except ImportError:
            print(f"[skip] python-docx missing, cannot read {path.name}", file=sys.stderr)
            return "", ext
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs), ext
    return "", ext


# ---------------------------------------------------------------------------
# Chapter detection
# ---------------------------------------------------------------------------


def _detect_chapters(text: str) -> tuple[list[str], dict[str, str]]:
    """Return (detected_canonical_slugs, coverage_map).

    coverage_map maps every canonical slug to ``present`` / ``partial`` /
    ``missing``.
    """
    headings = [m.group(1).strip().lower() for m in HEADING_RE.finditer(text or "")]
    text_lower = (text or "").lower()
    detected: list[str] = []
    coverage: dict[str, str] = {}
    for slug, keywords in CANONICAL_CHAPTERS:
        # Heading match → present; substring match anywhere → partial; else → missing
        in_heading = any(any(k in h for k in keywords) for h in headings)
        in_body = any(k in text_lower for k in keywords)
        if in_heading:
            coverage[slug] = "present"
            detected.append(slug)
        elif in_body:
            coverage[slug] = "partial"
        else:
            coverage[slug] = "missing"
    return detected, coverage


def _detect_language(text: str) -> str:
    if not text:
        return ""
    sample = text[:4000].lower()
    de_score = sum(
        sample.count(w)
        for w in [" und ", " der ", " die ", " das ", " ist ", " mit ", " für ", " nicht "]
    )
    en_score = sum(
        sample.count(w) for w in [" and ", " the ", " is ", " for ", " not ", " with ", " of "]
    )
    if de_score > en_score * 1.2:
        return "de"
    if en_score > de_score * 1.2:
        return "en"
    return ""


# ---------------------------------------------------------------------------
# Reusability candidate detection
# ---------------------------------------------------------------------------

REUSABILITY_PATTERNS = [
    (
        "loto-procedure",
        ["lockout", "tagout", "lockout/tagout", "loto", "gegen wiedereinschalten sichern"],
    ),
    (
        "signal-words-legend",
        ["danger", "warning", "caution", "notice", "gefahr", "warnung", "vorsicht", "hinweis"],
    ),
    (
        "target-groups-default",
        [
            "zielgruppe",
            "qualifikation",
            "target group",
            "qualification",
            "operator",
            "bediener",
            "elektrofachkraft",
        ],
    ),
    (
        "general-safety-rules",
        ["allgemeine sicherheitshinweise", "general safety", "betriebsanleitung lesen"],
    ),
    (
        "ppe-default",
        ["persönliche schutzausrüstung", "psa", "personal protective equipment", "ppe"],
    ),
    (
        "disposal-electronics-weee",
        ["weee", "elektroaltgeräte", "elektronikentsorgung", "disposal of electronic"],
    ),
    (
        "warning-structure-safe",
        ["safe-prinzip", "signalwort → art der gefahr → folgen", "signal word → hazard"],
    ),
]


def _find_reusability(files: list[FileInventory], texts: list[str]) -> list[dict]:
    """Find candidate content blocks that recur across ≥ 2 input files."""
    out: list[dict] = []
    for key, keywords in REUSABILITY_PATTERNS:
        hits = []
        for f, txt in zip(files, texts):
            t = (txt or "").lower()
            if any(k in t for k in keywords):
                hits.append(f.path)
        if len(hits) >= 2:
            out.append(
                {
                    "key": key,
                    "files_with_match": hits,
                    "match_count": len(hits),
                    "reusability": "high" if len(hits) >= len(files) * 0.6 else "medium",
                }
            )
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_dir", help="Directory of input manuals.")
    parser.add_argument("--out", default="-", help="Output JSON path (- = stdout).")
    parser.add_argument("--lang", default="", help="Override detected language.")
    parser.add_argument(
        "--max-files", type=int, default=50, help="Cap on number of files processed."
    )
    args = parser.parse_args()

    input_dir = Path(args.input_dir).resolve()
    if not input_dir.is_dir():
        print(f"error: {input_dir} is not a directory", file=sys.stderr)
        return 2

    files = sorted([p for p in input_dir.iterdir() if p.is_file()])[: args.max_files]

    file_invs: list[FileInventory] = []
    texts: list[str] = []
    for path in files:
        text, ext = _extract_text(path)
        lang = args.lang or _detect_language(text)
        chapters, coverage = _detect_chapters(text)
        file_invs.append(
            FileInventory(
                path=str(path.relative_to(input_dir)),
                bytes_size=path.stat().st_size,
                extension=ext,
                detected_language=lang,
                detected_chapters=chapters,
                coverage=coverage,
                text_excerpt=(text[:600] if text else ""),
            )
        )
        texts.append(text)

    chapter_freq: Counter[str] = Counter()
    for f in file_invs:
        chapter_freq.update(f.detected_chapters)

    inv = Inventory(
        input_dir=str(input_dir),
        file_count=len(file_invs),
        total_bytes=sum(f.bytes_size for f in file_invs),
        files=file_invs,
        canonical_chapters=[slug for slug, _ in CANONICAL_CHAPTERS],
        cross_file_chapter_frequency=dict(chapter_freq),
        reusability_candidates=_find_reusability(file_invs, texts),
    )

    out_dict = {
        "input_dir": inv.input_dir,
        "file_count": inv.file_count,
        "total_bytes": inv.total_bytes,
        "files": [asdict(f) for f in inv.files],
        "canonical_chapters": inv.canonical_chapters,
        "cross_file_chapter_frequency": inv.cross_file_chapter_frequency,
        "reusability_candidates": inv.reusability_candidates,
    }
    payload = json.dumps(out_dict, indent=2, ensure_ascii=False)
    if args.out == "-":
        print(payload)
    else:
        Path(args.out).write_text(payload, encoding="utf-8")
        print(f"wrote {args.out}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
